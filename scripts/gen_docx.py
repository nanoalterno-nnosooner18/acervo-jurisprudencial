# -*- coding: utf-8 -*-
"""Genera la versión Word .docx del acervo con índice jerárquico."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from acervo_dataset import ACERVO
from gen_html import materia_canonica, instancia_corta, construir_entradas_normalizadas


JERARQUIA_NOMBRES = {
    1: "Pleno de la Suprema Corte de Justicia de la Nación",
    2: "Salas de la Suprema Corte de Justicia de la Nación",
    3: "Plenos Regionales (antes Plenos de Circuito)",
    4: "Tribunales Colegiados de Circuito",
    5: "Tribunales Superiores de Justicia locales",
}

ACENTO = RGBColor(0x6B, 0x18, 0x18)
TINTA = RGBColor(0x1A, 0x1A, 0x1A)
TINTA_SUAVE = RGBColor(0x55, 0x55, 0x55)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def make_doc():
    doc = Document()

    # Estilo base
    style = doc.styles['Normal']
    style.font.name = 'Cambria'
    style.font.size = Pt(11)

    # márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.5)
        # numeración en footer
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(p)

    md = ACERVO["metadata"]

    # PORTADA
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(md["titulo"])
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = ACENTO

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(md["subtitulo"])
    r.font.size = Pt(13)
    r.font.italic = True
    r.font.color.rgb = TINTA_SUAVE

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Compilado: {md['fecha_compilacion']} · Versión {md['version']}").font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Fuente: {md['fuente']}").font.size = Pt(9)

    doc.add_paragraph()

    # Aviso
    aviso = doc.add_paragraph()
    aviso.paragraph_format.left_indent = Cm(1)
    aviso.paragraph_format.right_indent = Cm(1)
    r = aviso.add_run("ADVERTENCIA. ")
    r.bold = True; r.font.color.rgb = ACENTO; r.font.size = Pt(10)
    r = aviso.add_run(md["advertencia"])
    r.font.size = Pt(10); r.font.italic = True; r.font.color.rgb = TINTA_SUAVE

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    r = p.add_run("JERARQUÍA LEGAL. ")
    r.bold = True; r.font.color.rgb = ACENTO; r.font.size = Pt(10)
    r = p.add_run(md["jerarquia_legal"])
    r.font.size = Pt(10); r.font.italic = True; r.font.color.rgb = TINTA_SUAVE

    doc.add_page_break()

    # ÍNDICE
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("Í N D I C E   J E R Á R Q U I C O")
    r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = ACENTO

    doc.add_paragraph()

    entradas = construir_entradas_normalizadas()
    entradas.sort(key=lambda x: (x["jerarquia"], x["materia_principal"], x["tema"], x["id"]))

    # construir índice por niveles y materias
    grupos = {}
    for e in entradas:
        grupos.setdefault(e["jerarquia"], {}).setdefault(e["materia_principal"], {}).setdefault(e["tema"], []).append(e)

    for nivel in sorted(grupos.keys()):
        p = doc.add_paragraph()
        r = p.add_run(f"{nivel}. {JERARQUIA_NOMBRES[nivel]}")
        r.font.bold = True; r.font.size = Pt(13); r.font.color.rgb = ACENTO
        for mat in sorted(grupos[nivel].keys()):
            count_mat = sum(len(v) for v in grupos[nivel][mat].values())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            r = p.add_run(f"   · {mat}")
            r.font.size = Pt(11); r.font.bold = True
            r2 = p.add_run(f"  ({count_mat})")
            r2.font.size = Pt(10); r2.font.color.rgb = TINTA_SUAVE
            for tema in sorted(grupos[nivel][mat].keys()):
                count_t = len(grupos[nivel][mat][tema])
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.4)
                r = p.add_run(f"      › {tema}")
                r.font.size = Pt(10); r.font.color.rgb = TINTA_SUAVE
                r2 = p.add_run(f"  [{count_t}]")
                r2.font.size = Pt(9); r2.font.color.rgb = TINTA_SUAVE
        doc.add_paragraph()

    # Estadística
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Total de entradas: {len(entradas)}")
    r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = TINTA_SUAVE

    doc.add_page_break()

    # CONTENIDO POR NIVEL
    for nivel in sorted(grupos.keys()):
        # Encabezado nivel
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"NIVEL {nivel} — {JERARQUIA_NOMBRES[nivel].upper()}")
        r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = ACENTO

        doc.add_paragraph()

        for mat in sorted(grupos[nivel].keys()):
            p = doc.add_paragraph()
            r = p.add_run(f"Materia: {mat}")
            r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = ACENTO
            for tema in sorted(grupos[nivel][mat].keys()):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.3)
                r = p.add_run(f"› {tema}")
                r.font.size = Pt(12); r.font.bold = True; r.font.italic = True
                r.font.color.rgb = TINTA

                for e in grupos[nivel][mat][tema]:
                    # tabla por entrada (1 columna, varias filas) para que se mantenga unida en página
                    tbl = doc.add_table(rows=0, cols=1)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    tbl.autofit = False
                    # ancho
                    for col in tbl.columns:
                        col.width = Cm(15.5)

                    # fila: metadatos badges en texto
                    row = tbl.add_row().cells[0]
                    set_cell_shading(row, "F0E6D5")
                    p1 = row.paragraphs[0]
                    p1.paragraph_format.space_after = Pt(2)
                    r = p1.add_run(f"[{e['tipo'].upper()}]  ")
                    r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = ACENTO
                    r = p1.add_run(f"{e['instancia']}  ·  {e['epoca']}  ·  {e['ambito']}")
                    r.font.size = Pt(9); r.font.color.rgb = TINTA_SUAVE

                    # fila: rubro
                    row = tbl.add_row().cells[0]
                    p1 = row.paragraphs[0]
                    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    r = p1.add_run(f"Rubro. ")
                    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = ACENTO
                    r = p1.add_run(e["rubro"])
                    r.font.size = Pt(11); r.font.bold = True

                    # fila: meta localización
                    row = tbl.add_row().cells[0]
                    p1 = row.paragraphs[0]
                    bits = []
                    bits.append(f"Subtema: {e['subtema']}" if e.get("subtema") else "")
                    bits.append(f"Registro digital: {e['registro_digital']}")
                    if e.get("tesis_clave"):
                        bits.append(f"Tesis: {e['tesis_clave']}")
                    if e.get("tipo_resolucion"):
                        bits.append(f"Resolución: {e['tipo_resolucion']}")
                    if e.get("fecha_publicacion"):
                        bits.append(f"Fecha: {e['fecha_publicacion']}")
                    if e.get("fuente_publicacion"):
                        bits.append(f"Fuente: {e['fuente_publicacion']}")
                    if e.get("votacion"):
                        bits.append(f"Votación: {e['votacion']}")
                    bits.append(f"Vigencia: {e['vigencia']}")
                    txt = " · ".join([b for b in bits if b])
                    r = p1.add_run(txt)
                    r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = TINTA_SUAVE

                    # Fila extra: URL SJF
                    if e.get("url_sjf"):
                        row = tbl.add_row().cells[0]
                        p1 = row.paragraphs[0]
                        r = p1.add_run("Consulta SJF/SCJN: ")
                        r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = ACENTO
                        r = p1.add_run(e["url_sjf"])
                        r.font.size = Pt(9); r.font.color.rgb = ACENTO; r.font.underline = True

                    # fila: texto
                    row = tbl.add_row().cells[0]
                    p1 = row.paragraphs[0]
                    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    r = p1.add_run("Texto. ")
                    r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = ACENTO
                    r = p1.add_run(e["texto_resumen"])
                    r.font.size = Pt(10)

                    # fila: trascendencia
                    if e.get("trascendencia"):
                        row = tbl.add_row().cells[0]
                        p1 = row.paragraphs[0]
                        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        r = p1.add_run("Trascendencia. ")
                        r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = ACENTO
                        r = p1.add_run(e["trascendencia"])
                        r.font.size = Pt(10); r.font.italic = True

                    # fila: explicación
                    if e.get("explicacion"):
                        row = tbl.add_row().cells[0]
                        p1 = row.paragraphs[0]
                        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        r = p1.add_run("Explicación práctica. ")
                        r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = ACENTO
                        r = p1.add_run(e["explicacion"])
                        r.font.size = Pt(10)

                    # fila: etiquetas
                    if e.get("etiquetas"):
                        row = tbl.add_row().cells[0]
                        p1 = row.paragraphs[0]
                        r = p1.add_run("Etiquetas: ")
                        r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = TINTA_SUAVE
                        r = p1.add_run(" · ".join(f"#{t}" for t in e["etiquetas"]))
                        r.font.size = Pt(9); r.font.color.rgb = TINTA_SUAVE

                    # bordes simples
                    for row in tbl.rows:
                        for c in row.cells:
                            tcPr = c._tc.get_or_add_tcPr()
                            borders = OxmlElement('w:tcBorders')
                            for edge in ('top', 'left', 'bottom', 'right'):
                                b = OxmlElement(f'w:{edge}')
                                b.set(qn('w:val'), 'single')
                                b.set(qn('w:sz'), '4')
                                b.set(qn('w:color'), 'D8D2C4')
                                borders.append(b)
                            tcPr.append(borders)

                    doc.add_paragraph()

        if nivel < max(grupos.keys()):
            doc.add_page_break()

    # CIERRE
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— Fin del acervo —")
    r.font.italic = True; r.font.size = Pt(11); r.font.color.rgb = TINTA_SUAVE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Para consulta y verificación oficial: https://sjf2.scjn.gob.mx")
    r.font.size = Pt(9); r.font.color.rgb = TINTA_SUAVE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Acervo NaNo · Compilado por NanoLawyer (asistente jurídico mexicano)")
    r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = TINTA_SUAVE

    return doc


if __name__ == "__main__":
    out = Path("/sessions/brave-bold-goodall/mnt/outputs/acervo/Acervo_Jurisprudencial.docx")
    doc = make_doc()
    doc.save(out)
    print(f"DOCX generado: {out} ({out.stat().st_size:,} bytes)")
